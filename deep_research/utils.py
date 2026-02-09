import sys
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt
from google import genai
import asyncio
import re
from .config import GEMINI_KEY, GEMINI_MODEL, GEMINI_MAX_DELAY_S, GEMINI_MAX_CONCURRENCY, GEMINI_RPS, GEMINI_CIRCUIT_FAILURE_THRESHOLD, GEMINI_CIRCUIT_COOLDOWN_S
import time

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("deep_research.log"),
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger(__name__)


class ExternalServiceError(Exception):
    """Raised when an external service (LLM, API) fails after retries.

    Attributes:
        service: short name of the service (e.g., 'gemini')
        message: original error message
        status: optional status or code
    """
    def __init__(self, service: str, message: str, status: int = None):
        super().__init__(f"{service}: {message}")
        self.service = service
        self.message = message
        self.status = status

def log_error(context: str, error: str):
    """Log an error with context."""
    logger.error(f"[{context}] {error}")

def safe_save(doc: Document, base_path: Path):
    """Save document with a timestamp to avoid overwriting."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = base_path.with_name(f"{base_path.stem}_{timestamp}{base_path.suffix}")
    try:
        # Save to a temporary file in the same directory then atomically replace
        import io, os
        tmp_path = path.with_suffix(path.suffix + ".tmp")
        # Use an in-memory buffer then write bytes atomically to avoid partial files
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        tmp_path.parent.mkdir(parents=True, exist_ok=True)
        with open(tmp_path, "wb") as f:
            f.write(bio.read())
        os.replace(str(tmp_path), str(path))
        logger.info("✅ Saved report to %s", path)
    except Exception as e:
        logger.exception("Failed to save DOCX to %s: %s", path, e)

def build_doc(report: str) -> Document:
    """Convert markdown report to a Word document."""
    doc = Document()
    
    for line in report.splitlines():
        line = line.strip()
        if not line:
            continue
        
        if line == "---":
            doc.add_page_break()
            continue
        
        if line.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(line[2:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
            continue
        
        if line.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            continue
        
        if line.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(line[4:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            continue
        
        if line.startswith("#### "):
            p = doc.add_heading(level=4)
            run = p.add_run(line[4:])
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            continue
        
        para = doc.add_paragraph()
        
        # Handle bullet points
        if line.startswith("* "):
            para.style = 'List Bullet'
            line = line[2:]
        
        # Basic markdown parsing for bold and italic
        # Note: This is a simplified parser from the original code
        tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", line)
        
        for tok in tokens:
            if not tok:
                continue
                
            if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
                run = para.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
                run = para.add_run(tok[1:-1])
                run.italic = True
            else:
                # Fix common LLM markdown error: "*Title*:*" -> "*Title*:"
                # Only replace if it's plain text
                clean_tok = tok.replace(":*", ":")
                run = para.add_run(clean_tok)
            
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    
    return doc


def safe_write_text(path: Path, text: str, encoding: str = "utf-8") -> bool:
    """Write text to `path` atomically. Returns True on success."""
    try:
        tmp = path.with_name(f"{path.name}.tmp")
        # Ensure parent exists
        tmp.parent.mkdir(parents=True, exist_ok=True)
        tmp.write_text(text, encoding=encoding)
        # Atomic replace
        import os
        os.replace(str(tmp), str(path))
        logger.info("Wrote file %s (atomic)", path)
        return True
    except Exception as e:
        logger.exception("Failed to write file %s: %s", path, e)
        try:
            if tmp.exists():
                tmp.unlink()
        except Exception as inner_e:
            logger.exception("Failed to remove temp file %s: %s", tmp, inner_e)
        return False

# Initialize Gemini Client lazily
_client = None

# Concurrency and rate-limiting primitives for Gemini (per event loop)
_gemini_runtime = {}

# Circuit-breaker state
_gemini_fail_count = 0
_gemini_circuit_open_until = 0.0


def _parse_retry_delay_seconds(error_str: str) -> Optional[float]:
    """Extract retry delay seconds from Gemini error messages."""
    retry_match = re.search(r"retryDelay['\"]?:\s*['\"]?(\d+(?:\.\d+)?)s", error_str)
    if not retry_match:
        retry_match = re.search(r"retryDelay['\"]?:\s*['\"]?(\d+(?:\.\d+)?)ms", error_str)
        if retry_match:
            return float(retry_match.group(1)) / 1000.0
    if retry_match:
        return float(retry_match.group(1))

    retry_match = re.search(r"retry in\s+(\d+(?:\.\d+)?)s", error_str, re.IGNORECASE)
    if not retry_match:
        retry_match = re.search(r"retry in\s+(\d+(?:\.\d+)?)ms", error_str, re.IGNORECASE)
        if retry_match:
            return float(retry_match.group(1)) / 1000.0
    if retry_match:
        return float(retry_match.group(1))

    return None

def get_client():
    global _client
    if _client is None:
        if not GEMINI_KEY:
            raise ValueError("GEMINI_KEY not found in environment variables.")
        _client = genai.Client(api_key=GEMINI_KEY)
    return _client


def reset_client():
    """Reset the cached Gemini client (for API key rotation in UI)."""
    global _client
    _client = None
    logger.info("Gemini client reset")

async def gemini_complete(prompt: str, max_tokens: int = 6000) -> str:
    """Generate text using Gemini."""
    max_retries = 5
    backoff = 2
    # Respect circuit-breaker
    now = time.monotonic()
    global _gemini_fail_count, _gemini_circuit_open_until
    if now < _gemini_circuit_open_until:
        raise ExternalServiceError("gemini", f"circuit-open until {_gemini_circuit_open_until - now:.1f}s")

    # Enforce concurrency and RPS limits (per event loop)
    delay_s = 1.0 / max(1.0, GEMINI_RPS)

    for attempt in range(max_retries + 1):
        try:
            loop = asyncio.get_running_loop()
            rt = _gemini_runtime.get(id(loop))
            if rt is None:
                rt = {
                    "semaphore": asyncio.Semaphore(GEMINI_MAX_CONCURRENCY),
                    "lock": asyncio.Lock(),
                    "state": {"last_request": 0.0},
                }
                _gemini_runtime[id(loop)] = rt

            async with rt["semaphore"]:
                # throttle requests by RPS
                async with rt["lock"]:
                    now = time.monotonic()
                    wait_time = rt["state"].get("last_request", 0.0) + delay_s - now
                    if wait_time > 0:
                        await asyncio.sleep(wait_time)
                    rt["state"]["last_request"] = time.monotonic()

                client = get_client()
                response = client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=prompt,
                    config=genai.types.GenerateContentConfig(
                        max_output_tokens=max_tokens,
                        temperature=0.3
                    )
                )
                # Successful call: reset failure counter
                _gemini_fail_count = 0
                return response.text
        except Exception as e:
            # Immediately record exception so failures are visible to pre-mortem checks
            logger.exception("Gemini API exception: %s", e)
            import traceback
            error_str = str(e)
            # If rate-limited, backoff and retry
            if "503" in error_str or "429" in error_str:
                if attempt < max_retries:
                    retry_delay = _parse_retry_delay_seconds(error_str)
                    if retry_delay is not None and retry_delay > 0:
                        wait_time = min(retry_delay, GEMINI_MAX_DELAY_S)
                    else:
                        wait_time = min(backoff * (2 ** attempt), GEMINI_MAX_DELAY_S)
                    logger.warning("Gemini API error (%s). Retrying in %ss...", error_str, wait_time)
                    await asyncio.sleep(wait_time)
                    continue

            # Record failure for circuit-breaker
            _gemini_fail_count += 1
            if _gemini_fail_count >= GEMINI_CIRCUIT_FAILURE_THRESHOLD:
                _gemini_circuit_open_until = time.monotonic() + GEMINI_CIRCUIT_COOLDOWN_S
                logger.error("Gemini circuit opened for %ss due to repeated failures", GEMINI_CIRCUIT_COOLDOWN_S)

            logger.exception("Gemini API error: %s", error_str)
            # After exhausting retries, raise a structured error so callers can decide how to handle
            raise ExternalServiceError("gemini", error_str) from e
